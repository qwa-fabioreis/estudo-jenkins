from docx import Document
import numpy as np
import logging


logger = logging.getLogger()

class Croqui:
    d_sections = {
        'PAINEL FRONTAL SUPERIOR': None,
        'PAINEL FRONTAL CENTRAL': None,
        'PAINEL FRONTAL INFERIOR': None,
        'LATERAL ESQUERDA': None,
        'LATERAL DIREITA': None,
        'BASE': None,
        'PAINEL VERSO SUPERIOR': None,
        'PAINEL VERSO CENTRAL': None,
        'PAINEL VERSO INFERIOR': None
    }

    def __init__(self, doc):
        self.read_docx(doc)

    class SectionContent:
        rows = []

        def __init__(self):
            self.rows = []

        class Row:
            texts = []
            tables = []

            def __init__(self, texts, tables):
                self.texts = texts
                self.tables = tables

            def __str__(self):
                text = '\n'.join([str(row) for row in self.texts])
                text += '\n'.join([str(table) for table in self.tables])

            class Table:
                rows = []
                is_processed = False

                def __init__(self, rows):
                    self.rows = rows

                def __str__(self):
                    return '\n'.join([str(cell) for cell in [row for row in self.rows]])

            class Paragraph:
                text_contents = []

                def __init__(self, text_contents):
                    self.text_contents = text_contents

            class TextContent:
                text = ''
                found = False
                processed = False

                def __init__(self, text):
                    self.text = text

                def __str__(self):
                    return self.text

    def read_docx(self, doc_path):
        logger.info('[Croqui] starting to process docx')
        
        document = Document(doc_path)
        for table in document.tables:

            title = table.rows[0].cells[1].paragraphs[0].text.upper().strip()

            if title in self.d_sections:

                self.d_sections[title] = Croqui.SectionContent()
                for row in table.rows[1:]:
                    row_data = self.extract_data_from_row(row)
                    if row_data is not None:
                        self.d_sections[title].rows.append(row_data)

    def extract_data_from_row(self, row):

        if len(row.cells) >= 3:
            paragraphs = self.extract_text(row)
            tables = []
            if self.verify_if_row_contains_text(row, 'enriquecimento por quilograma'):
                paragraphs.append(self.convert_enriquecimento_por_kg_to_text(row))
            else:
                tables = self.extract_table(row)
            row_content = Croqui.SectionContent.Row(paragraphs, tables)

            return row_content

    def extract_text(self, row):

        paragraphs = []

        for paragraph in row.cells[1].paragraphs:

            textos = []
            
            for palavra in paragraph.text.strip(' ').replace('–', '-').split('. '):
                # for palavra_split_virgula in palavra.split(', '):
                textos.append(palavra)

            # Ira armanzenar todos os textos na class TextContent para formar um objeto da classe Paragraph
            TextContent = []

            for texto in textos:
                if len(texto) >= 4:
                    TextContent.append(Croqui.SectionContent.Row.TextContent(texto))

            paragraphs.append(Croqui.SectionContent.Row.Paragraph(TextContent))

        return paragraphs




    def extract_table(self, row):
        tables = []
        for table_detail in row.cells[1].tables:

            table_text = []
            columns_len = len(table_detail.rows[0].cells)
            if self.verify_if_row_contains_text(row, 'Níveis de garantia'):
                table_text.append([Croqui.SectionContent.Row.TextContent('Níveis de garantia')] +
                                  [Croqui.SectionContent.Row.TextContent('')]*(columns_len - 1))

            if self.verify_if_row_contains_text(row, 'Recomendação diária de consumo*'):
                table_text.append([Croqui.SectionContent.Row.TextContent('Recomendação diária de consumo*')] +
                                  [Croqui.SectionContent.Row.TextContent('')]*(columns_len - 1))

            for row_detail in table_detail.rows:
                line = []
                for cell_detail in row_detail.cells:
                    if len(cell_detail.paragraphs) > 0 and len(cell_detail.paragraphs[0].text) > 0:
                        line.append(Croqui.SectionContent.Row.TextContent(cell_detail
                                                                          .paragraphs[0].text.replace('\xa0', '').replace('–','-')))
                    else:
                        line.append(Croqui.SectionContent.Row.TextContent(' '))

                table_text.append(line)

            tables.append(Croqui.SectionContent.Row.Table(np.array(table_text)))


        return tables


    def convert_enriquecimento_por_kg_to_text(self, row):
        text = ''

        for table_detail in row.cells[1].tables:

            for row_detail in table_detail.rows:
                text += row_detail.cells[0].text.lower() + ' (' + row_detail.cells[1].text + ' ' + row_detail.cells[
                    2].text + '), '

        paragraph = Croqui.SectionContent.Row.Paragraph(
            [Croqui.SectionContent.Row.TextContent(text.strip(','))]
        )
        return paragraph


    def verify_if_row_contains_text(self, row, text):
        found = False
        for para in row.cells[1].paragraphs:
            if text.upper() in para.text.upper():
                found = True
                break

        return found
