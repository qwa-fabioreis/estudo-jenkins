import os
import docx
import zipfile

from io import BytesIO
from fastapi import UploadFile, File
from celery.utils.log import get_task_logger


logger = get_task_logger(__name__)

def croqui_image(croqui_doc: UploadFile = File(...)) -> bytes:
    """
    Croqui Image - v2.0

    Este metodo recebe um arquivo docx da API, encontra a imagem da segunda tabela do
    croqui que corresponde a um animal que estampa a embalagem, e ira retornar
    os bytes da imagem encontrada.

    @param croqui_doc: UploadFile
    @return: imageBytes: bytes
    """

    logger.info('[croqui_image] starting to process file')
    # Lendo os bytes do arquivo recebido pela API
    document_bytes = BytesIO(croqui_doc.file.read())

    # Abertura do documento
    document = docx.Document(document_bytes)

    # Abertura do documento como zip em modo leitura, apenas para retirada da imagem
    documentZip = zipfile.ZipFile(document_bytes, 'r')

    # Ira guardar o diretorio da imagem contida no zip
    imagedir_zip: str

    # Dicionario para guardar os rId das imagens e seus respectivos nomes
    rId_imagens: dict = {}

    # Busca por todos os rIds dentro do arquivo docx
    for r in document.part.rels.values():
        if isinstance(r._target, docx.parts.image.ImagePart):
            rId_imagens[r.rId] = os.path.basename(r._target.partname)

    # Buscando pela imagem contida na segunda tabela do documento
    for paragraph in document.tables[1].rows[0].cells[1].paragraphs:

        if 'graphic' in paragraph._p.xml:
            for rId in rId_imagens:
                if rId in paragraph._p.xml:
                    imagedir_zip = 'word/media/' + str(rId_imagens[rId])

    return documentZip.read(imagedir_zip)
