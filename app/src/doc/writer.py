from docx import Document
from docx.shared import Inches, RGBColor, Cm
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from tempfile import TemporaryFile

from app.src.file.file_utils import FileUtils
from ..config.config import Config
import cv2
from celery.utils.log import get_task_logger

logger = get_task_logger(__name__)
CORRECT_STYLE = 'correct'
INCORRECT_STYLE = 'incorrect'


def write_document(crop: str, croqui, d_imgs, input_directory: str):

    logger.info('[write_document] starting to write final report')

    document = Document()

    document.add_heading('Resultados', 0)

    add_style(document, CORRECT_STYLE, RGBColor(0x0, 0xb0, 0x0))
    add_style(document, INCORRECT_STYLE, RGBColor(0xb0, 0x0, 0x0))

    for section_name in croqui.d_sections:
        logger.info('Writing section %s' % section_name)

        table = document.add_table(rows=0, cols=2)
        col1 = table.columns[0]
        col1.width = Cm(10)
        col2 = table.columns[1]
        col2.width = Cm(5)

        hdr_cells = table.add_row().cells
        orange_shade = parse_xml(
            r'<w:shd {} w:fill="EC711F"/>'.format(nsdecls('w')))
        hdr_cells[0]._tc.get_or_add_tcPr().append(orange_shade)
        hdr_cells[0].bold = True
        orange_shade = parse_xml(
            r'<w:shd {} w:fill="EC711F"/>'.format(nsdecls('w')))
        hdr_cells[1]._tc.get_or_add_tcPr().append(orange_shade)
        hdr_cells[0].add_paragraph(section_name)
        hdr_cells[1].text = ''
        if croqui.d_sections[section_name] is not None:
            for doc_row in croqui.d_sections[section_name].rows:
                outer_row_cells = table.add_row().cells
                outer_row_cells[1].text = ''

                for paragraph in doc_row.texts:

                    for text_content in paragraph.text_contents:
                        outer_row_cells[0].add_paragraph(text_content.text,
                                                         style=get_correct_or_incorrect_style(text_content.found))

                for table_content in doc_row.tables:

                    table1 = outer_row_cells[0].add_table(
                        rows=0, cols=len(table_content.rows[-1]))

                    for i, row in enumerate(table_content.rows):
                        row_cells = table1.add_row().cells

                        for j, cell in enumerate(row):
                            row_cells[j].paragraphs[0].text = cell.text
                            row_cells[j].paragraphs[0].style = get_correct_or_incorrect_style(
                                cell.found)

                        if i == 0:
                            merge_cells(row, row_cells)

        add_img(document, d_imgs[section_name])
        document.add_page_break()

    document.add_page_break()

    logger.info('[write_document] saving final report')

    save_report(document, crop, input_directory)
    move_report_to_success(input_directory)


def move_report_to_success(input_directory: str):
    source = Config.in_progress_folder() + input_directory

    FileUtils.move_directory(source, Config.success_folder())


def save_report(document: Document, crop: str, input_directory: str):
    timestamp_as_date_string = Config.transform_dir_into_date(input_directory)
    report_name = get_report_name(crop, timestamp_as_date_string)
    report_directory = get_report_directory(input_directory, crop)

    FileUtils.save_document(document, report_directory, report_name)


def get_report_name(crop: str, timestamp_as_date_string: str) -> str:
    return "/report_" + crop + "_" + timestamp_as_date_string + ".docx"


def get_report_directory(input_directory: str, crop: str) -> str:
    return Config.in_progress_folder() + input_directory + "/" + crop


def merge_cells(row, docx_row_cells):
    title = row[0].text
    for l, cell in enumerate(row[1:]):
        if title != cell.text:
            break
    if l > 0:

        a = docx_row_cells[0]
        b = docx_row_cells[l+1]
        A = a.merge(b)
        for i in range(0, len(A.paragraphs)-1):
            delete_paragraph(A.paragraphs[0])


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


def add_style(document, name, color):
    style = document.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    font = style.font
    font.color.rgb = color


def get_correct_or_incorrect_style(bool):
    return CORRECT_STYLE if bool else INCORRECT_STYLE


def add_img(document, img):
    with TemporaryFile() as temp_image:
        success, buffer = cv2.imencode(".png", img)
        buffer.tofile(temp_image)
        temp_image.seek(0)
        document.add_picture(temp_image, width=Inches(7))
